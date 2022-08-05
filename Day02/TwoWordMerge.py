from docx import Document
from pathlib import Path
def mrge_without_format(docfiles:list):
    '''
    只获取内容进行合并
    '''
    #遍历文件
    for docx_file in sorted(docfiles):
        another_doc = Document(docx_file)
        # 获取每个文件的的所有段落
        paras = another_doc.paragraphs
        for para in paras:
            # 为新的word文件创建一个新段落
            newpar = doc.add_paragraph('')
            #将提取的内容写入新的文本段落中
            newpar.add_run(para.text)

        #所有文件合并完成后在指定路径中保存
        doc.save(Path(word_files_path,'new.docx'))

    #调用函数
    merge-mrge_without_format(files)