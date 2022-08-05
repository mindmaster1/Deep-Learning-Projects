from docx import Document
from docx.shared import RGBColor
from pathlib import Path, PurePath

# word文件所在路径
word_files_path = 'D:/Python自动化办公/Day02/word样例文件'

# # 取得该目录下所有的docx格式文件
# p = Path(word_files_path)

my_content = '''
野老篱边江岸回，柴门不正逐江开。
渔人网集澄潭下，贾客船随返照来。
长路关心悲剑阁，片云何意傍琴台？
王师未报收东郡，城阙秋生画角哀。
'''

def add_content_mode1(content):

    '''
    增加内容
    '''
    print(doc)
    para = doc.add_paragraph().add_run(content)

    # 设置字体格式
    para.font.name = '仿宋'
    # 设置下划线
    para.font.underline = True
    # 设置颜色
    para.font.color.rgb = RGBColor(255, 128, 128)
    # 其他设置参考官方文档
    # https://python-docx.readthedocs.io/en/latest/api/text.html#run-objects


doc = Document()
add_content_mode1(my_content)
doc.save(Path(word_files_path, 'new2.docx'))
